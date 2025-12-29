import os
import re
import requests
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from youtube_transcript_api import (
    YouTubeTranscriptApi,
    TranscriptsDisabled,
    NoTranscriptFound,
    VideoUnavailable
)
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from textwrap import wrap 
from docx import Document
from io import BytesIO
import unicodedata



load_dotenv()


GROQ_API_KEY = os.getenv("GROQ_API_KEY")

app = Flask(__name__)
CORS(app)

PROMPT = """
You are a YouTube video summarizer.
Summarize the transcript below into a clean, concise bullet-point summary
within 250 words. Use clear, structured formatting.

Transcript:
"""



def clean_text(text: str)->str:
    # normalize unicode 
    text = unicodedata.normalize("NFKC", text)
    
    # remove space b/w single letters 
    text = re.sub(r'(?<=w)\s(?=\w)', '', text)
    
    # restore space afte punctuation
    text = re.sub(r'([.,!?])(\w)', r'\1 \2', text)
    
    # fixed bullet spacing
    text = re.sub(r'-\s*','- ',text)
    
    return text



@app.route('/')
def index():
    return render_template('index.html')



def extract_transcript_details(youtube_video_url: str) -> str:
    try:
        match = re.search(r"(?:v=|youtu\.be/)([a-zA-Z0-9_-]{11})", youtube_video_url)
        if not match:
            raise ValueError("Invalid YouTube URL format")

        video_id = match.group(1)

        transcript_items = YouTubeTranscriptApi().fetch(video_id)

        transcript = " ".join(item.text for item in transcript_items)

        if not transcript.strip():
            raise Exception("Transcript is empty")

        return transcript[:5000]

    except TranscriptsDisabled:
        raise Exception("Transcripts are disabled for this video")

    except NoTranscriptFound:
        raise Exception("No transcript available for this video")

    except VideoUnavailable:
        raise Exception("Video is unavailable or private")

    except Exception as e:
        raise Exception(f"Transcript fetch failed: {str(e)}")



def generate_summary(transcript: str) -> str:
    url = "https://api.groq.com/openai/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "openai/gpt-oss-120b",
        "messages": [
            {"role": "system", "content": "You are a YouTube video summarizer."},
            {"role": "user", "content": PROMPT + transcript}
        ],
        "temperature": 0.7
    }

    response = requests.post(url, headers=headers, json=payload, timeout=30)

    if response.status_code != 200:
        raise Exception(f"Groq API Error: {response.status_code} {response.text}")

    return response.json()["choices"][0]["message"]["content"]



@app.route('/api/download_txt', methods=['POST'])
def download_txt():
    data = request.get_json()
    summary = data.get("summary", "")

    buffer = BytesIO()
    buffer.write(summary.encode("utf-8"))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="video-summary.txt",
        mimetype="text/plain"
    )
    
    

@app.route('/api/summarize', methods=['POST'])
def summarize_video():
    data = request.get_json(silent=True)
    video_url = data.get("url") if data else None

    if not video_url:
        return jsonify({"error": "No YouTube URL provided"}), 400

    try:
        transcript = extract_transcript_details(video_url)
        raw_summary = generate_summary(transcript)
        summary = clean_text(raw_summary)
        return jsonify({"summary": summary}), 200

    except Exception as e:
        print("ERROR:", str(e))
        return jsonify({"error": str(e)}), 500




@app.route('/api/download_doc', methods=['POST'])
def download_doc():
    data = request.get_json()
    summary = clean_text(data.get('summary', ''))

    doc = Document()
    doc.add_heading("Video Summary", level=1)

    for line in summary.split("\n"):
        if line.strip().startswith('-'):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)
            
            
    f = BytesIO()
    doc.save(f)
    f.seek(0)

    return send_file(
        f,
        as_attachment=True,
        download_name="video-summary.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port, debug=True)